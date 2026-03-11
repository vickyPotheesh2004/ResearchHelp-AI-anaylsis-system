import os
os.environ["LOKY_MAX_CPU_COUNT"] = "1"
import concurrent.futures
import streamlit as st
import streamlit.components.v1 as components
import io
import re
import json
import os
from datetime import datetime
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from src.extractor import process_file
from src.qa_engine import QAEngine
from src.config import validate_config
from src.logging_utils import setup_logging

# Initialize logging at application startup
setup_logging()

# Maximum chat input length to prevent abuse
MAX_CHAT_INPUT_LENGTH = 2000


def sanitize_filename(filename: str) -> str:
    """Sanitize uploaded filename to prevent path traversal attacks."""
    # Remove any path components
    filename = os.path.basename(filename)
    # Remove any special characters except alphanumeric, underscore, hyphen, and dot
    filename = re.sub(r'[^\w\s.-]', '', filename)
    # Limit filename length
    max_length = 200
    if len(filename) > max_length:
        name, ext = os.path.splitext(filename)
        filename = name[:max_length - len(ext)] + ext
    return filename or "unnamed_file"


def sanitize_for_markdown(text: str) -> str:
    """Sanitize user-provided text to prevent HTML/Markdown injection."""
    if not text:
        return ""
    # Escape HTML special characters
    text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    # Escape markdown special characters that could break rendering
    text = text.replace('[', '&#91;').replace(']', '&#93;')
    return text


def render_content_with_mermaid(content):
    import urllib.parse
    
    image_pattern = re.compile(r'<image_prompt>(.*?)</image_prompt>', re.DOTALL)
    images = image_pattern.findall(content)
    clean_content = image_pattern.sub('', content)

    mermaid_pattern = re.compile(r'```mermaid\s*\n(.*?)```', re.DOTALL)
    parts = mermaid_pattern.split(clean_content)

    if len(parts) == 1:
        st.markdown(clean_content)
    else:
        for i, part in enumerate(parts):
            if i % 2 == 0:
                text = part.strip()
                if text:
                    st.markdown(text)
            else:
                mermaid_code = part.strip().replace('`', "'") # Escape backticks that would break JS string
                html = f"""
                <div style="background:#ffffff; border:1px solid #e0e0e0; border-radius:12px; padding:20px; margin:12px 0; text-align:center;">
                    <div class="mermaid" style="display:flex; justify-content:center; min-height: 200px;">
                        {mermaid_code}
                    </div>
                </div>
                <script src="https://cdn.jsdelivr.net/npm/mermaid@11.4.0/dist/mermaid.min.js"></script>
                <script>
                    try {{
                        mermaid.initialize({{
                            startOnLoad: true,
                            theme: 'default',
                            flowchart: {{ curve: 'basis', padding: 20 }},
                            securityLevel: 'loose',
                            suppressErrorIndicators: false
                        }});
                    }} catch (e) {{
                        console.error('Mermaid init error:', e);
                    }}
                </script>
                """
                components.html(html, height=500, scrolling=True)

    for img_prompt in images:
        encoded = urllib.parse.quote(img_prompt.strip())
        image_url = f"https://image.pollinations.ai/prompt/{encoded}?width=800&height=400&nologo=true"
        st.image(image_url, caption="AI Generated Visualization", use_container_width=True)


def _clean_for_speech(text):
    text = re.sub(r'<image_prompt>.*?</image_prompt>', 'Image generated.', text, flags=re.DOTALL)
    text = re.sub(r'```mermaid.*?```', 'Diagram included.', text, flags=re.DOTALL)
    text = re.sub(r'```.*?```', '', text, flags=re.DOTALL)
    text = re.sub(r'#{1,6}\s*', '', text)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'`(.*?)`', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    text = re.sub(r'[|\-]{3,}', '', text)
    text = re.sub(r'<[^>]+>', '', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def render_voice_controller():
    """Injects a self-healing persistent Microphone icon into the chat bar."""
    
    html = f"""
    <div style="display:none">
        <script>
            // Persistent state
            var recognition = null;
            var isRecording = false;
            
            const WinParent = window.parent;
            if (WinParent && (WinParent.webkitSpeechRecognition || WinParent.SpeechRecognition)) {{
                const SpeechRecognition = WinParent.SpeechRecognition || WinParent.webkitSpeechRecognition;
                recognition = new SpeechRecognition();
                recognition.continuous = false;
                recognition.interimResults = false;
                recognition.lang = 'en-US';
                
                recognition.onstart = function() {{ isRecording = true; updateBtnUI(true); }};
                recognition.onend = function() {{ isRecording = false; updateBtnUI(false); }};
                recognition.onerror = function() {{ isRecording = false; updateBtnUI(false); }};
                recognition.onresult = function(event) {{ 
                    injectToChat(event.results[0][0].transcript); 
                }};
            }}
            
            function updateBtnUI(recording) {{
                WinParent.document.querySelectorAll('.voice-btn-mic').forEach(btn => {{
                    if(recording) btn.classList.add('recording');
                    else btn.classList.remove('recording');
                }});
            }}
            
            function injectToChat(text) {{
                const ta = WinParent.document.querySelector('textarea[data-testid="stChatInputTextArea"]');
                if (ta) {{
                    const nativeSetter = Object.getOwnPropertyDescriptor(window.HTMLTextAreaElement.prototype, "value").set;
                    nativeSetter.call(ta, text);
                    ta.dispatchEvent(new Event('input', {{ bubbles: true }}));
                    ta.dispatchEvent(new Event('change', {{ bubbles: true }}));
                    setTimeout(() => {{
                        const submitBtn = WinParent.document.querySelector('button[data-testid="stChatInputSubmitButton"]');
                        if (submitBtn) submitBtn.click();
                    }}, 300);
                }}
            }}
            
            function toggleMic() {{
                if (!recognition) {{ alert("Voice input not supported in this browser."); return; }}
                if (isRecording) recognition.stop();
                else recognition.start();
            }}

            function injectButtons() {{
                // Target the chat bar or its parent container
                const chatBar = WinParent.document.querySelector('[data-testid="stChatInput"]');
                if (!chatBar) return;

                // Check if the button is MISSING from this specific chatBar instance
                if (!chatBar.querySelector('.voice-btn-mic')) {{
                    console.log("ResearchHelp-AI-anaylsis-system: Injecting self-healing mic...");
                    
                    const container = WinParent.document.createElement('div');
                    container.className = 'voice-btn-container';
                    
                    const micBtn = WinParent.document.createElement('div');
                    micBtn.className = 'voice-btn voice-btn-mic';
                    micBtn.innerHTML = '🎙️';
                    micBtn.onclick = toggleMic;
                    
                    container.appendChild(micBtn);
                    chatBar.appendChild(container);
                }}
            }}
            
            // High-frequency polling (Every 200ms) to ensure mic never disappears
            if (WinParent) {{
                clearInterval(WinParent.docmindMicInterval);
                WinParent.docmindMicInterval = setInterval(injectButtons, 200);
                injectButtons();
                console.log("ResearchHelp-AI-anaylsis-system: Voice controller self-healing active.");
            }}
        </script>
    </div>
    """
    components.html(html, height=0, width=0)


def speak_text(text, key):
    clean = _clean_for_speech(text)
    safe = clean.replace('\\', '\\\\').replace('`', '').replace('"', '&quot;').replace("'", "\\'").replace('\n', ' ').replace('\r', '')
    # Truncate for very long content to avoid browser limits
    if len(safe) > 5000:
        safe = safe[:5000] + '... Content truncated for audio playback.'
    html = f"""
    <div style="margin: 4px 0;">
        <button id="tts-btn-{key}" onclick="toggleSpeech_{key}()" style="
            background: linear-gradient(135deg, #302b63, #24243e);
            color: #a8edea; border: 1px solid #4a45a0; border-radius: 8px;
            padding: 6px 16px; cursor: pointer; font-size: 0.8rem;
            font-weight: 600; font-family: Inter, sans-serif;
            transition: all 0.2s ease;
        " onmouseover="this.style.background='linear-gradient(135deg,#4a45a0,#302b63)'"
           onmouseout="this.style.background='linear-gradient(135deg,#302b63,#24243e)'">
            🔊 Listen
        </button>
    </div>
    <script>
        var speaking_{key} = false;
        var utterance_{key} = null;
        function toggleSpeech_{key}() {{
            var btn = document.getElementById('tts-btn-{key}');
            if (speaking_{key}) {{
                window.speechSynthesis.cancel();
                speaking_{key} = false;
                btn.innerHTML = '🔊 Listen';
                btn.style.borderColor = '#4a45a0';
            }} else {{
                window.speechSynthesis.cancel();
                utterance_{key} = new SpeechSynthesisUtterance('{safe}');
                utterance_{key}.rate = 1.0;
                utterance_{key}.pitch = 1.0;
                utterance_{key}.onend = function() {{
                    speaking_{key} = false;
                    btn.innerHTML = '🔊 Listen';
                    btn.style.borderColor = '#4a45a0';
                }};
                window.speechSynthesis.speak(utterance_{key});
                speaking_{key} = true;
                btn.innerHTML = '⏹️ Stop';
                btn.style.borderColor = '#a8edea';
            }}
        }}
    </script>
    """
    components.html(html, height=42)

st.set_page_config(
    page_title="AI Document Q&A System",
    page_icon="🔬",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Validate configuration on startup
is_valid, config_errors = validate_config()
if not is_valid:
    st.error("⚠️ Configuration Error:")
    for error in config_errors:
        st.error(f"  • {error}")
    st.info("Please set up your .env file with a valid OpenRouter API key. Copy .env.example to .env and add your key.")
    st.stop()

CUSTOM_CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');

[data-testid="stAppViewContainer"] {
    font-family: 'Inter', sans-serif;
    background-color: #0f0c29;
}

/* Voice Interaction UI */
[data-testid="stChatInput"] {
    position: relative;
    padding-right: 100px !important;
}

.voice-btn-container {
    position: absolute;
    right: 12px;
    bottom: 12px;
    display: flex;
    gap: 8px;
    z-index: 1000;
}

.voice-btn {
    width: 36px;
    height: 36px;
    border-radius: 50%;
    background: linear-gradient(135deg, #1e3c72, #2a5298);
    display: flex;
    align-items: center;
    justify-content: center;
    cursor: pointer;
    box-shadow: 0 4px 12px rgba(0,0,0,0.3);
    border: 1px solid rgba(255,255,255,0.1);
    transition: all 0.2s cubic-bezier(0.4, 0, 0.2, 1);
    font-size: 1.1rem;
}

.voice-btn:hover {
    transform: translateY(-2px);
    box-shadow: 0 6px 16px rgba(118, 75, 162, 0.4);
}

.voice-btn:active {
    transform: scale(0.9);
}

.voice-btn.recording {
    background: linear-gradient(135deg, #ff416c, #ff4b2b);
    animation: pulse-red 1.5s infinite;
}

@keyframes pulse-red {
    0% { box-shadow: 0 0 0 0 rgba(255, 65, 108, 0.7); }
    70% { box-shadow: 0 0 0 10px rgba(255, 65, 108, 0); }
    100% { box-shadow: 0 0 0 0 rgba(255, 65, 108, 0); }
}

/* Fix for st.tabs rendering and visibility */
.stTabs [data-baseweb="tab-list"] {
    gap: 10px;
}
.stTabs [data-baseweb="tab"] {
    border-radius: 8px 8px 0 0;
    padding: 10px 20px;
    background-color: #1e1e2e;
}
.stTabs [aria-selected="true"] {
    background-color: #302b63 !important;
}

.main-header {
    background: linear-gradient(135deg, #0f0c29, #302b63, #24243e);
    padding: 1.8rem 2rem;
    border-radius: 16px;
    margin-bottom: 1.5rem;
    color: white;
    text-align: center;
    box-shadow: 0 8px 32px rgba(48, 43, 99, 0.3);
}
.main-header h1 {
    margin: 0; font-size: 1.8rem; font-weight: 700;
    background: linear-gradient(90deg, #a8edea, #fed6e3);
    -webkit-background-clip: text; -webkit-text-fill-color: transparent;
}
.main-header p {
    margin: 0.4rem 0 0; font-size: 0.9rem; opacity: 0.8; color: #c0c0d0;
}

.intent-badge {
    display: inline-flex; align-items: center; gap: 6px;
    padding: 4px 14px; border-radius: 20px; font-size: 0.78rem;
    font-weight: 600; margin-bottom: 8px;
}
.intent-document_qa { background: #1a3a5c; color: #7ec8f0; border: 1px solid #2a5a8c; }
.intent-suggestion_request { background: #3a2a1a; color: #f0c87e; border: 1px solid #6a4a2a; }
.intent-research_addon { background: #1a3a2a; color: #7ef0a8; border: 1px solid #2a6a4a; }
.intent-off_topic { background: #3a1a1a; color: #f07e7e; border: 1px solid #6a2a2a; }

.source-card {
    background: #1e1e2e; border: 1px solid #313147; border-radius: 10px;
    padding: 12px 16px; margin-bottom: 8px; font-size: 0.82rem;
}
.source-card .source-header {
    display: flex; justify-content: space-between; margin-bottom: 4px;
}
.source-card .source-file { color: #a8edea; font-weight: 600; }
.source-card .source-topic { color: #fed6e3; font-size: 0.75rem; }
.source-card .source-score {
    background: #302b63; color: #a8edea; padding: 2px 8px;
    border-radius: 10px; font-size: 0.7rem; font-weight: 600;
}
.source-card .source-preview { color: #888; font-size: 0.78rem; margin-top: 4px; }

.stats-grid {
    display: grid; grid-template-columns: 1fr 1fr; gap: 8px; margin: 10px 0;
}
.stat-card {
    background: linear-gradient(135deg, #1e1e2e, #2a2a3e);
    border: 1px solid #313147; border-radius: 10px;
    padding: 12px; text-align: center;
}
.stat-card .stat-value {
    font-size: 1.5rem; font-weight: 700;
    background: linear-gradient(90deg, #a8edea, #fed6e3);
    -webkit-background-clip: text; -webkit-text-fill-color: transparent;
}
.stat-card .stat-label { font-size: 0.72rem; color: #888; margin-top: 2px; }

.suggestion-card {
    background: #1e1e2e; border: 1px solid #313147; border-radius: 10px;
    padding: 12px 16px; margin-bottom: 8px; cursor: pointer;
    transition: all 0.2s ease;
}
.suggestion-card:hover {
    border-color: #a8edea; transform: translateY(-1px);
    box-shadow: 0 4px 12px rgba(168, 237, 234, 0.1);
}
.suggestion-title { color: #e0e0f0; font-weight: 600; font-size: 0.85rem; }
.suggestion-desc { color: #888; font-size: 0.78rem; margin-top: 4px; }
.suggestion-cat {
    display: inline-block; padding: 2px 8px; border-radius: 10px;
    font-size: 0.68rem; font-weight: 600; margin-top: 6px;
}
.cat-improvement { background: #1a3a5c; color: #7ec8f0; }
.cat-innovation { background: #3a1a3a; color: #c87ef0; }
.cat-gap { background: #3a2a1a; color: #f0c87e; }
.cat-research { background: #1a3a2a; color: #7ef0a8; }
.cat-optimization { background: #2a2a1a; color: #e0e07e; }

.topic-chip {
    display: inline-block; background: #1e1e2e; border: 1px solid #313147;
    border-radius: 16px; padding: 4px 12px; margin: 3px 4px 3px 0;
    font-size: 0.75rem; color: #c0c0d0;
}

.stDownloadButton > button {
    background: linear-gradient(135deg, #302b63, #24243e) !important;
    color: white !important; border: 1px solid #4a45a0 !important;
    border-radius: 10px !important; font-weight: 600 !important;
    transition: all 0.3s ease !important;
}
.stDownloadButton > button:hover {
    background: linear-gradient(135deg, #4a45a0, #302b63) !important;
    box-shadow: 0 4px 12px rgba(74, 69, 160, 0.3) !important;
}
</style>
"""

st.markdown(CUSTOM_CSS, unsafe_allow_html=True)

st.markdown("""
<div class="main-header">
    <h1>🔬 ResearchHelp-AI-anaylsis-system: AI Document Q&A System</h1>
    <p>Upload documents • Ask questions • Get research insights • Download analysis</p>
</div>
""", unsafe_allow_html=True)

if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "topics_found" not in st.session_state:
    st.session_state.topics_found = {}
if "system_ready" not in st.session_state:
    st.session_state.system_ready = False
if "ai_engine" not in st.session_state:
    st.session_state.ai_engine = QAEngine()
if "auto_suggestions" not in st.session_state:
    st.session_state.auto_suggestions = []
if "doc_overview" not in st.session_state:
    st.session_state.doc_overview = ""
if "ieee_metadata" not in st.session_state:
    st.session_state.ieee_metadata = {
        "title": "",
        "authors": "",
        "emails": "",
        "colleges": "",
        "additional_notes": ""
    }
# Lock for concurrent upload protection
if "processing_lock" not in st.session_state:
    st.session_state.processing_lock = False


def generate_markdown_export(history, overview="", suggestions=None, stats=None):
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    md = f"# 🔬 ResearchHelp-AI-anaylsis-system: AI Document Q&A System Report\n\n"
    md += f"**Generated**: {timestamp}\n\n---\n\n"

    if stats:
        md += "## 📊 Session Analytics\n\n"
        md += f"| Metric | Value |\n|---|---|\n"
        md += f"| Questions Asked | {stats.get('questions_asked', 0)} |\n"
        md += f"| Topics Explored | {stats.get('topics_accessed', 0)} |\n"
        md += f"| Sources Referenced | {stats.get('sources_used', 0)} |\n"
        md += f"| Total Topics Available | {stats.get('total_topics', 0)} |\n\n"

    if overview:
        md += f"## 📄 Document Overview\n\n{overview}\n\n---\n\n"

    if suggestions:
        md += "## 💡 AI-Generated Suggestions\n\n"
        for i, s in enumerate(suggestions, 1):
            md += f"### {i}. {s.get('title', 'Suggestion')}\n"
            md += f"{s.get('description', '')}\n"
            md += f"*Category: {s.get('category', 'general')}*\n\n"
        md += "---\n\n"

    md += "## 💬 Research Session Log\n\n"
    for msg in history:
        role = "👤 Researcher" if msg["role"] == "user" else "🧠 AI Analyst"
        md += f"### {role}\n"

        if msg.get("intent"):
            intent = msg["intent"]
            md += f"*Intent: {intent.get('emoji', '')} {intent.get('label', '')}*\n\n"

        md += f"{msg['content']}\n\n"

        if msg.get("reasoning_details") and msg["reasoning_details"] != "Model processed logic internally (Invisible Reasoning Pipeline).":
            md += f"<details><summary>Logic Trace</summary>\n\n> {msg['reasoning_details']}\n\n</details>\n\n"

        if msg.get("sources"):
            md += "**Sources Referenced:**\n"
            for src in msg["sources"]:
                md += f"- [{src['file']}] {src['topic']} (relevance: {src['score']})\n"
            md += "\n"

        md += "---\n\n"

    return md


def generate_html_export(history, overview="", suggestions=None, stats=None):
    html = """<!DOCTYPE html>
<html><head><meta charset="utf-8">
<title>ResearchHelp-AI-anaylsis-system: AI Document Q&A System Report</title>
<style>
body { font-family: 'Segoe UI', sans-serif; max-width: 900px; margin: 0 auto; padding: 40px; background: #0f0c29; color: #e0e0f0; }
h1 { background: linear-gradient(90deg, #a8edea, #fed6e3); -webkit-background-clip: text; -webkit-text-fill-color: transparent; text-align: center; }
h2 { color: #a8edea; border-bottom: 1px solid #313147; padding-bottom: 8px; }
.msg { background: #1e1e2e; border: 1px solid #313147; border-radius: 12px; padding: 16px 20px; margin: 12px 0; }
.user { border-left: 3px solid #a8edea; }
.assistant { border-left: 3px solid #fed6e3; }
.role { font-weight: 700; font-size: 0.9rem; margin-bottom: 8px; }
.intent-tag { display: inline-block; padding: 2px 10px; border-radius: 12px; font-size: 0.75rem; background: #302b63; color: #a8edea; margin-bottom: 8px; }
.stat-grid { display: grid; grid-template-columns: repeat(4, 1fr); gap: 12px; margin: 16px 0; }
.stat { background: #1e1e2e; border: 1px solid #313147; border-radius: 10px; padding: 16px; text-align: center; }
.stat-val { font-size: 1.6rem; font-weight: 700; color: #a8edea; }
.stat-lbl { font-size: 0.75rem; color: #888; }
table { width: 100%; border-collapse: collapse; margin: 12px 0; }
th, td { border: 1px solid #313147; padding: 8px 12px; text-align: left; }
th { background: #1e1e2e; color: #a8edea; }
hr { border: none; border-top: 1px solid #313147; margin: 24px 0; }
</style></head><body>
"""
    html += f"<h1>🔬 ResearchHelp-AI-anaylsis-system: AI Document Q&A System Report</h1>"
    html += f"<p style='text-align:center;color:#888;'>Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p><hr>"

    if stats:
        html += "<h2>📊 Session Analytics</h2><div class='stat-grid'>"
        for key, label in [("questions_asked", "Questions"), ("topics_accessed", "Topics"), ("sources_used", "Sources"), ("total_topics", "Total Topics")]:
            html += f"<div class='stat'><div class='stat-val'>{stats.get(key, 0)}</div><div class='stat-lbl'>{label}</div></div>"
        html += "</div>"

    if overview:
        html += f"<h2>📄 Document Overview</h2><div class='msg'>{overview}</div>"

    html += "<h2>💬 Research Session</h2>"
    for msg in history:
        role_class = "user" if msg["role"] == "user" else "assistant"
        role_label = "👤 Researcher" if msg["role"] == "user" else "🧠 AI Analyst"
        html += f"<div class='msg {role_class}'><div class='role'>{role_label}</div>"
        if msg.get("intent"):
            html += f"<span class='intent-tag'>{msg['intent'].get('emoji', '')} {msg['intent'].get('label', '')}</span>"
        content = msg['content'].replace('\n', '<br>')
        html += f"<div>{content}</div></div>"

    html += "</body></html>"
    return html


def generate_docx_export(history, overview="", suggestions=None, stats=None):
    doc = DocxDocument()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    title = doc.add_heading('ResearchHelp-AI-anaylsis-system: AI Document Q&A System Report', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    if stats:
        doc.add_heading('Session Analytics', level=1)
        table = doc.add_table(rows=1, cols=4, style='Light Grid Accent 1')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        hdr[0].text = 'Questions Asked'
        hdr[1].text = 'Topics Explored'
        hdr[2].text = 'Sources Used'
        hdr[3].text = 'Total Topics'
        row = table.add_row().cells
        row[0].text = str(stats.get('questions_asked', 0))
        row[1].text = str(stats.get('topics_accessed', 0))
        row[2].text = str(stats.get('sources_used', 0))
        row[3].text = str(stats.get('total_topics', 0))
        for cell in row:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(14)
        doc.add_paragraph()

    if overview:
        doc.add_heading('Document Overview', level=1)
        for line in overview.split('\n'):
            line = line.strip()
            if not line:
                continue
            if line.startswith('## '):
                doc.add_heading(line.replace('## ', '').replace('#', ''), level=2)
            elif line.startswith('### '):
                doc.add_heading(line.replace('### ', '').replace('#', ''), level=3)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                clean = line.replace('**', '').replace('*', '').replace('`', '')
                doc.add_paragraph(clean)

    if suggestions:
        doc.add_heading('AI-Generated Suggestions', level=1)
        for i, s in enumerate(suggestions, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {s.get('title', f'Suggestion {i}')}")
            run.font.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)
            doc.add_paragraph(s.get('description', ''))
            cat = doc.add_paragraph()
            cat_run = cat.add_run(f"Category: {s.get('category', 'general').upper()}")
            cat_run.font.italic = True
            cat_run.font.size = Pt(9)
            cat_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_heading('Research Q&A Session', level=1)
    for msg in history:
        role = "Researcher (User)" if msg["role"] == "user" else "AI Analyst"
        p = doc.add_paragraph()
        role_run = p.add_run(f"{role}")
        role_run.font.bold = True
        role_run.font.size = Pt(11)
        if msg["role"] == "user":
            role_run.font.color.rgb = RGBColor(0x1a, 0x5a, 0x8c)
        else:
            role_run.font.color.rgb = RGBColor(0x8c, 0x1a, 0x5a)

        if msg.get("intent") and msg["role"] == "assistant":
            intent = msg["intent"]
            intent_p = doc.add_paragraph()
            intent_run = intent_p.add_run(f"[{intent.get('emoji', '')} {intent.get('label', '')}]")
            intent_run.font.italic = True
            intent_run.font.size = Pt(9)
            intent_run.font.color.rgb = RGBColor(0x66, 0x66, 0x99)

        content = msg['content']
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                continue
            if line.startswith('## '):
                doc.add_heading(line.replace('## ', '').replace('#', ''), level=3)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                clean = line.replace('**', '').replace('*', '').replace('`', '')
                doc.add_paragraph(clean)

        if msg.get("sources") and msg["role"] == "assistant":
            src_p = doc.add_paragraph()
            src_run = src_p.add_run("Sources: ")
            src_run.font.bold = True
            src_run.font.size = Pt(9)
            for src in msg["sources"]:
                src_run = src_p.add_run(f"[{src['file']} | {src['topic']}] ")
                src_run.font.size = Pt(9)
                src_run.font.color.rgb = RGBColor(0x44, 0x44, 0x88)

        doc.add_paragraph('─' * 60).runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_ieee_docx(content, metadata):
    doc = DocxDocument()
    
    # IEEE papers usually have a specific style. 
    # We'll simulate a professional academic layout.
    
    # Title
    title_text = metadata.get("title") or "Research Paper Title"
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(title_text.upper())
    run_title.font.bold = True
    run_title.font.size = Pt(24)
    run_title.font.name = 'Times New Roman'
    
    # Authors
    p_authors = doc.add_paragraph()
    p_authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_authors = p_authors.add_run(metadata.get("authors") or "Author Names Not Provided")
    run_authors.font.size = Pt(11)
    run_authors.font.name = 'Times New Roman'
    
    # Affiliations & Emails
    p_affil = doc.add_paragraph()
    p_affil.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affil_text = f"{metadata.get('colleges', '')}\n{metadata.get('emails', '')}"
    run_affil = p_affil.add_run(affil_text)
    run_affil.font.italic = True
    run_affil.font.size = Pt(10)
    run_affil.font.name = 'Times New Roman'
    
    doc.add_paragraph() # Spacer
    
    # Split content into sections
    sections = content.split('###')
    for section in sections:
        section = section.strip()
        if not section:
            continue
            
        lines = section.split('\n')
        header = lines[0].strip()
        body = '\n'.join(lines[1:]).strip()
        
        # Section Heading
        p_head = doc.add_paragraph()
        run_head = p_head.add_run(header.upper())
        run_head.font.bold = True
        run_head.font.size = Pt(12)
        run_head.font.name = 'Times New Roman'
        
        # Section Body
        for b_line in body.split('\n'):
            b_line = b_line.strip()
            if not b_line: continue
            
            p_body = doc.add_paragraph()
            p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Remove markdown bold/italics for the DOCX
            clean_line = b_line.replace('**', '').replace('*', '').replace('`', '')
            
            # Simple bullet point detection
            if b_line.startswith('- ') or b_line.startswith('* '):
                p_body.style = 'List Bullet'
                run_body = p_body.add_run(clean_line[2:])
            else:
                run_body = p_body.add_run(clean_line)
                
            run_body.font.size = Pt(10)
            run_body.font.name = 'Times New Roman'

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# ─── SIDEBAR ───

with st.sidebar:
    st.markdown("### 📁 Document Upload")
    uploaded_files = st.file_uploader(
        "Select files to process",
        accept_multiple_files=True,
        type=["pdf", "docx", "txt", "csv", "xlsx", "png", "jpg", "jpeg"]
    )

    if st.button("🚀 Process Documents", use_container_width=True):
        if uploaded_files:
            # Check for concurrent processing lock
            if st.session_state.processing_lock:
                st.warning("⚠️ Processing in progress. Please wait for the current operation to complete.")
            else:
                # Acquire lock
                st.session_state.processing_lock = True
                try:
                    # Sanitize filenames to prevent path traversal attacks
                    data = {sanitize_filename(f.name): process_file(f)[1] for f in uploaded_files}
                    st.session_state.topics_found = {}
                    st.session_state.chat_history = []
                    st.session_state.auto_suggestions = []
                    st.session_state.doc_overview = ""
                    
                    # Reset IEEE metadata conditionally or keep it? 
                    # Usually keep it as users might upload new files for the same paper

                    status_text = st.empty()

                    def update_ui(fn, cur, total, topic):
                        if fn not in st.session_state.topics_found:
                            st.session_state.topics_found[fn] = []
                        st.session_state.topics_found[fn].append(topic)
                        status_text.text(f"Processing: {fn}\nIdentified: {topic}")

                    with st.spinner("Extracting text and running semantic segmentation..."):
                        try:
                            st.session_state.ai_engine.ingest_and_segment(data, progress_callback=update_ui)
                        except Exception as e:
                            st.error(f"Error processing documents: {str(e)}")
                            st.session_state.processing_lock = False
                            st.rerun()

                    with st.spinner("Generating document insights..."):
                        try:
                            chunks, metas = st.session_state.ai_engine.get_all_chunks()
                            if chunks:
                                st.session_state.doc_overview = st.session_state.ai_engine.research_engine.generate_document_overview(chunks, metas)
                                st.session_state.auto_suggestions = st.session_state.ai_engine.research_engine.generate_auto_suggestions(chunks, metas)
                        except Exception as e:
                            st.warning(f"Could not generate insights: {str(e)}")

                    st.session_state.system_ready = True
                    st.success("✅ Processing complete!")
                except Exception as e:
                    st.error(f"Error processing documents: {str(e)}")
                finally:
                    # Release lock
                    st.session_state.processing_lock = False
        else:
            st.warning("Please upload at least one file.")

    st.divider()
    if st.session_state.topics_found:
        st.divider()
        st.markdown("### 🗂️ Identified Topics")
        for filename, topics in st.session_state.topics_found.items():
            st.markdown(f"**{filename}**")
            unique_topics = list(dict.fromkeys(topics))
            topic_html = "".join([f'<span class="topic-chip">{t}</span>' for t in unique_topics])
            st.markdown(topic_html, unsafe_allow_html=True)

    if st.session_state.system_ready:
        st.divider()
        stats = st.session_state.ai_engine.get_session_stats()
        st.markdown("### 📊 Session Analytics")
        st.markdown(f"""
        <div class="stats-grid">
            <div class="stat-card"><div class="stat-value">{stats['questions_asked']}</div><div class="stat-label">Questions</div></div>
            <div class="stat-card"><div class="stat-value">{stats['topics_accessed']}</div><div class="stat-label">Topics Hit</div></div>
            <div class="stat-card"><div class="stat-value">{stats['sources_used']}</div><div class="stat-label">Sources</div></div>
            <div class="stat-card"><div class="stat-value">{stats['total_topics']}</div><div class="stat-label">Total Topics</div></div>
        </div>
        """, unsafe_allow_html=True)

    if st.session_state.chat_history:
        st.divider()
        st.markdown("### 📥 Export Session")

        stats = st.session_state.ai_engine.get_session_stats()
        col1, col2, col3 = st.columns(3)

        with col1:
            st.download_button(
                label="📝 Markdown",
                data=generate_markdown_export(
                    st.session_state.chat_history,
                    st.session_state.doc_overview,
                    st.session_state.auto_suggestions,
                    stats
                ),
                file_name=f"research_report_{datetime.now().strftime('%Y%m%d_%H%M')}.md",
                mime="text/markdown",
                use_container_width=True
            )

        with col2:
            st.download_button(
                label="🌐 HTML",
                data=generate_html_export(
                    st.session_state.chat_history,
                    st.session_state.doc_overview,
                    st.session_state.auto_suggestions,
                    stats
                ),
                file_name=f"research_report_{datetime.now().strftime('%Y%m%d_%H%M')}.html",
                mime="text/html",
                use_container_width=True
            )

        with col3:
            st.download_button(
                label="📄 Word",
                data=generate_docx_export(
                    st.session_state.chat_history,
                    st.session_state.doc_overview,
                    st.session_state.auto_suggestions,
                    stats
                ),
                file_name=f"research_report_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
    if not st.session_state.get("system_ready", False):
        st.markdown("""
        <div style="text-align: center; padding: 40px 20px;">
            <h1 style="color: #a8edea; margin-bottom: 0;">ResearchHelp-AI-anaylsis-system AI</h1>
            <p style="color: #666;">Document Research & Voice Lab</p>
        </div>
        """, unsafe_allow_html=True)
    
    render_voice_controller()
    
    if st.session_state.system_ready: # Changed from uploaded_file to system_ready to match existing logic
        st.divider()
        if st.button("🗑️ Clear Session", use_container_width=True):
            st.session_state.chat_history = []
            st.session_state.topics_found = {}
            st.session_state.system_ready = False
            st.session_state.auto_suggestions = []
            st.session_state.doc_overview = ""
            st.rerun()


# ─── MAIN CONTENT ───

if not st.session_state.system_ready:
    st.markdown("""
    <div style="text-align: center; padding: 60px 20px;">
        <h2 style="color: #a8edea; font-weight: 600;">Welcome to the Research Lab 🧪</h2>
        <p style="color: #888; font-size: 1.1rem; max-width: 600px; margin: 0 auto;">
            Upload your documents in the sidebar to begin. The system will automatically
            segment topics, generate insights, and prepare for your research questions.
        </p>
        <div style="margin-top: 30px; display: grid; grid-template-columns: repeat(4, 1fr); gap: 16px; max-width: 700px; margin-left: auto; margin-right: auto;">
            <div style="background: #1e1e2e; border: 1px solid #313147; border-radius: 12px; padding: 20px; text-align: center;">
                <div style="font-size: 1.8rem;">📄</div>
                <div style="color: #a8edea; font-size: 0.85rem; font-weight: 600; margin-top: 8px;">Document Q&A</div>
                <div style="color: #666; font-size: 0.72rem; margin-top: 4px;">Ask any question</div>
            </div>
            <div style="background: #1e1e2e; border: 1px solid #313147; border-radius: 12px; padding: 20px; text-align: center;">
                <div style="font-size: 1.8rem;">💡</div>
                <div style="color: #f0c87e; font-size: 0.85rem; font-weight: 600; margin-top: 8px;">Suggestions</div>
                <div style="color: #666; font-size: 0.72rem; margin-top: 4px;">Get improvements</div>
            </div>
            <div style="background: #1e1e2e; border: 1px solid #313147; border-radius: 12px; padding: 20px; text-align: center;">
                <div style="font-size: 1.8rem;">🔬</div>
                <div style="color: #7ef0a8; font-size: 0.85rem; font-weight: 600; margin-top: 8px;">Research</div>
                <div style="color: #666; font-size: 0.72rem; margin-top: 4px;">Propose add-ons</div>
            </div>
            <div style="background: #1e1e2e; border: 1px solid #313147; border-radius: 12px; padding: 20px; text-align: center;">
                <div style="font-size: 1.8rem;">📥</div>
                <div style="color: #fed6e3; font-size: 0.85rem; font-weight: 600; margin-top: 8px;">Export</div>
                <div style="color: #666; font-size: 0.72rem; margin-top: 4px;">Download reports</div>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)
else:
    # ─── Document Overview Tab & Suggestions Tab & IEEE Metadata Tab & Chat Tab ───
    if st.session_state.doc_overview or st.session_state.auto_suggestions:
        tab_overview, tab_suggestions, tab_ieee, tab_chat = st.tabs(["📄 Document Overview", "💡 AI Suggestions", "📝 IEEE Metadata", "💬 Document Chat"])

        with tab_overview:
            if st.session_state.doc_overview:
                speak_text(st.session_state.doc_overview, "overview")
                render_content_with_mermaid(st.session_state.doc_overview)
            else:
                st.info("No overview generated yet.")

        with tab_suggestions:
            if st.session_state.auto_suggestions:
                st.markdown("### 💡 AI-Generated Research Suggestions")
                st.caption("These suggestions were auto-generated by analyzing your documents.")
                for i, s in enumerate(st.session_state.auto_suggestions):
                    cat = s.get("category", "general")
                    cat_class = f"cat-{cat}" if cat in ["improvement", "innovation", "gap", "research", "optimization"] else "cat-improvement"
                    st.markdown(f"**{i+1}. {s.get('title', f'Suggestion {i+1}')}**")
                    st.markdown(f"{s.get('description', '')}")
                    st.markdown(f"*Category: {cat.upper()}*")
                    st.divider()
            else:
                st.info("No suggestions generated yet.")

            if st.session_state.auto_suggestions:
                all_sug_text = " ".join([f"{s.get('title','')}: {s.get('description','')}" for s in st.session_state.auto_suggestions])
                speak_text(all_sug_text, "suggestions")

        with tab_ieee:
            st.markdown("### 📝 IEEE Paper Metadata")
            st.caption("Provide information for the official IEEE paper generation.")
            st.session_state.ieee_metadata["title"] = st.text_input("Paper Title", st.session_state.ieee_metadata["title"], key="ieee_title")
            st.session_state.ieee_metadata["authors"] = st.text_area("Team Members (Names)", st.session_state.ieee_metadata["authors"], placeholder="John Doe, Jane Smith...", key="ieee_authors")
            st.session_state.ieee_metadata["emails"] = st.text_area("Official Emails", st.session_state.ieee_metadata["emails"], placeholder="john@college.edu, jane@college.edu", key="ieee_emails")
            st.session_state.ieee_metadata["colleges"] = st.text_area("Colleges / Organizations", st.session_state.ieee_metadata["colleges"], key="ieee_colleges")
            st.session_state.ieee_metadata["additional_notes"] = st.text_area("Additional Project Notes", st.session_state.ieee_metadata["additional_notes"], key="ieee_notes")
            
            if st.checkbox("Ready for Generation"):
                st.info("You can now ask the AI to 'Generate an IEEE paper' in the Document Chat tab.")

    else:
        tab_chat = st.container()
        tab_overview = None
        tab_suggestions = None
        tab_ieee = None

    # ─── Chat Interface ───
    with tab_chat if (st.session_state.doc_overview or st.session_state.auto_suggestions) else st.container():
        # Render existing history FIRST so input stays at the bottom
        for msg in st.session_state.chat_history:
            with st.chat_message(msg["role"]):
                if msg.get("intent") and msg["role"] == "assistant":
                    intent = msg["intent"]
                    intent_class = f"intent-{intent.get('intent', 'document_qa')}"
                    st.markdown(
                        f'<div class="intent-badge {intent_class}">{intent.get("emoji", "📄")} {intent.get("label", "Document Q&A")}</div>',
                        unsafe_allow_html=True
                    )
                if msg.get("reasoning_details") and msg["role"] == "assistant":
                    with st.expander("🧠 View Logic Trace"):
                        st.text(msg["reasoning_details"])
                render_content_with_mermaid(msg["content"])
                if msg["role"] == "assistant":
                    speak_text(msg["content"], f"hist_{id(msg)}")
                if msg.get("sources") and msg["role"] == "assistant":
                    with st.expander(f"📚 Sources Referenced ({len(msg['sources'])})"):
                        for src in msg["sources"]:
                            # Sanitize user-provided source data
                            safe_file = sanitize_for_markdown(src.get('file', ''))
                            safe_topic = sanitize_for_markdown(src.get('topic', ''))
                            safe_preview = sanitize_for_markdown(src.get('preview', ''))
                            st.markdown(f"""
                            <div class="source-card">
                                <div class="source-header">
                                    <span class="source-file">{safe_file}</span>
                                    <span class="source-score">{src['score']}</span>
                                </div>
                                <div class="source-topic">{safe_topic}</div>
                                <div class="source-preview">{safe_preview}</div>
                            </div>
                            """, unsafe_allow_html=True)
                
                # IEEE Download Button (History)
                if msg["role"] == "assistant" and msg.get("intent", {}).get("intent") == "ieee_paper_gen":
                    ieee_docx = generate_ieee_docx(msg["content"], st.session_state.ieee_metadata)
                    st.download_button(
                        label="📄 Download IEEE Official Paper (.docx)",
                        data=ieee_docx,
                        file_name=f"IEEE_Paper_{st.session_state.ieee_metadata.get('title', 'Generated').replace(' ', '_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"dl_ieee_{id(msg)}",
                        use_container_width=True
                    )

        user_q = st.chat_input("Ask about your documents, request suggestions, or propose research add-ons...")
        if user_q:
            # Validate input length
            if len(user_q) > MAX_CHAT_INPUT_LENGTH:
                st.error(f"Input too long. Maximum {MAX_CHAT_INPUT_LENGTH} characters allowed.")
                st.rerun()
            if not user_q.strip():
                st.warning("Please enter a valid question.")
                st.rerun()
            
            st.session_state.chat_history.append({"role": "user", "content": user_q})
            
            with st.chat_message("user"):
                st.markdown(user_q)
                
            with st.chat_message("assistant"):
                stream_meta = {}
                final_data = {}
                response_placeholder = st.empty()
                streamed_text = ""

                # Pass IEEE metadata to the engine
                for event in st.session_state.ai_engine.get_answer_stream(
                    user_q, 
                    st.session_state.chat_history[:-1],
                    metadata=st.session_state.ieee_metadata
                ):
                    if event["type"] == "meta":
                        stream_meta = event
                        intent = event.get("intent", {})
                        intent_class = f"intent-{intent.get('intent', 'document_qa')}"
                        st.markdown(
                            f'<div class="intent-badge {intent_class}">{intent.get("emoji", "📄")} {intent.get("label", "Document Q&A")}</div>',
                            unsafe_allow_html=True
                        )
                    elif event["type"] == "token":
                        streamed_text += event["token"]
                        response_placeholder.markdown(streamed_text + "▌")
                    elif event["type"] == "done":
                        final_data = event

                final_content = final_data.get("content", streamed_text)
                response_placeholder.empty()
                render_content_with_mermaid(final_content)
                speak_text(final_content, "stream_latest")

                reasoning = final_data.get("reasoning")
                if reasoning:
                    with st.expander("🧠 View Logic Trace"):
                        st.text(reasoning)

                sources = stream_meta.get("sources", [])
                if sources:
                    with st.expander(f"📚 Sources Referenced ({len(sources)})"):
                        for src in sources:
                            # Sanitize user-provided source data
                            safe_file = sanitize_for_markdown(src.get('file', ''))
                            safe_topic = sanitize_for_markdown(src.get('topic', ''))
                            safe_preview = sanitize_for_markdown(src.get('preview', ''))
                            st.markdown(f"""
                            <div class="source-card">
                                <div class="source-header">
                                    <span class="source-file">{safe_file}</span>
                                    <span class="source-score">{src['score']}</span>
                                </div>
                                <div class="source-topic">{safe_topic}</div>
                                <div class="source-preview">{safe_preview}</div>
                            </div>
                            """, unsafe_allow_html=True)

                # IEEE Download Button (Current)
                if stream_meta.get("intent", {}).get("intent") == "ieee_paper_gen":
                    ieee_docx = generate_ieee_docx(final_content, st.session_state.ieee_metadata)
                    st.download_button(
                        label="📄 Download IEEE Official Paper (.docx)",
                        data=ieee_docx,
                        file_name=f"IEEE_Paper_{st.session_state.ieee_metadata.get('title', 'Generated').replace(' ', '_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="dl_ieee_latest",
                        use_container_width=True
                    )

            st.session_state.chat_history.append({
                "role": "assistant",
                "content": final_content,
                "reasoning_details": final_data.get("reasoning"),
                "intent": stream_meta.get("intent"),
                "sources": stream_meta.get("sources", []),
            })
            st.rerun()
